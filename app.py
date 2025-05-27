from fastapi import FastAPI, HTTPException
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field
from typing import Callable, Literal, Optional
import uvicorn
from rag_agent.chains.interview_chain import (
    get_assessment_chain,
    get_evaluate_chain,
    get_followup_chain,
    get_interview_chain,
    get_model_answer_chain,
)
import logging
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import OpenAIEmbeddings
from langchain.text_splitter import CharacterTextSplitter
from langchain.docstore.document import Document
from uuid import uuid4

import os
import io
import shutil

# PDF/DOCX 파싱
import PyPDF2
import docx

# 로깅 설정
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(title="AI Interview Simulator")
dist_path = os.path.join(os.path.dirname(__file__), "frontend/dist")

# 메모리 컨텍스트 저장 변수
stored_resume: Optional[str] = None
stored_jd: Optional[str] = None
# 사전 계산된 회사 정보 및 체인 저장 변수
stored_company_info: Optional[str] = None
interview_chain = None
followup_chain = None
evaluate_chain = None
model_answer_chain = None
assessment_chain = None
base_chain_inputs: Optional[dict] = None
# RAG 벡터 스토어
vectorstore: Optional[Chroma] = None
# 영속 디렉토리 설정 (환경변수 또는 기본 경로)
persist_directory = os.getenv(
    "CHROMA_DB_PATH",
    os.path.join(os.path.dirname(__file__), "rag_agent/vectorstore/chroma_db"),
)

ContentType = Literal["question", "answer", "modelAnswer", "evaluate"]
SpeakerType = Literal["agent", "user"]


class ChatItem(BaseModel):
    id: str
    type: ContentType  # "question" | "answer" | "modelAnswer"
    speaker: SpeakerType  # "agent" | "user"
    content: str


class ChatHistory(BaseModel):
    history: list[ChatItem] = []

    def add(self, type: ContentType, speaker: SpeakerType, content: str) -> str:
        id = str(uuid4().hex[:8])
        self.history.append(
            ChatItem(id=id, type=type, speaker=speaker, content=content)
        )
        return id

    def get_all_history(self) -> list[ChatItem]:
        return self.history

    def get_latest_question_id(self) -> Optional[str]:
        for item in reversed(self.history):
            if item.type == "question":
                return item.id
        return None

    def get_question_by_id(self, question_id: str) -> Optional[ChatItem]:
        for item in self.history:
            if item.id == question_id and item.type == "question":
                return item
        return None

    def validate_question_exists(self, question_id: str) -> bool:
        return any(
            item.id == question_id and item.type == "question" for item in self.history
        )

    def get_all_history_as_string(self) -> str:
        return "\n".join(
            [
                f"{item.speaker}: {item.content}"
                for item in self.history
                if item.type == "question" or item.type == "answer"
            ]
        )

    # def add_question(self, question: str):
    #     self.question_history.append(
    #         {"question_id": len(self.question_history), "question": question}
    #     )

    #     return self.question_history[-1]["question_id"]

    # def add_answer(self, question_id: str, answer: str):
    #     self.answer_history.append({"question_id": question_id, "answer": answer})

    #     return self.answer_history[-1]["question_id"]


class AnswerRequest(BaseModel):
    questionId: str
    content: str


class Question(BaseModel):
    id: str
    content: str


class Answer(BaseModel):
    id: str
    content: str


chat_history = ChatHistory()


# 로컬 파일 시스템에서 context와 회사 자료 자동 로딩
def parse_file_to_text(file_path: str) -> str:
    with open(file_path, "rb") as f:
        content = f.read()
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        if file_path.lower().endswith(".pdf"):
            reader = PyPDF2.PdfReader(io.BytesIO(content))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        elif file_path.lower().endswith((".docx", ".doc")):
            doc = docx.Document(io.BytesIO(content))
            return "\n".join(p.text for p in doc.paragraphs)
        else:
            return content.decode("utf-8", errors="ignore")


async def init_local_data():
    """로컬 db를 모두 비운다"""
    global vectorstore
    if os.path.exists(persist_directory):
        shutil.rmtree(persist_directory)
    os.makedirs(persist_directory, exist_ok=True)
    vectorstore = None
    logger.info("Vectorstore reset successfully.")


@app.on_event("startup")
async def load_local_data():
    await init_local_data()
    global stored_resume, stored_jd, vectorstore, stored_company_info, interview_chain, followup_chain, base_chain_inputs, chat_history, evaluate_chain, model_answer_chain, assessment_chain
    base_dir = os.path.join(os.path.dirname(__file__), "data")
    # 이력서 로딩
    resume_dir = os.path.join(base_dir, "resume")
    for fname in os.listdir(resume_dir):
        stored_resume = parse_file_to_text(os.path.join(resume_dir, fname))
        break
    # JD 로딩
    jd_dir = os.path.join(base_dir, "jd")
    for fname in os.listdir(jd_dir):
        stored_jd = parse_file_to_text(os.path.join(jd_dir, fname))
        break
    # 회사 자료 로딩 및 인덱싱
    company_dir = os.path.join(base_dir, "company_infos")
    docs = []
    for fname in os.listdir(company_dir):
        text = parse_file_to_text(os.path.join(company_dir, fname))
        splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
        for chunk in splitter.split_text(text):
            docs.append(Document(page_content=chunk, metadata={"filename": fname}))
    embeddings = OpenAIEmbeddings()
    vectorstore = Chroma(
        persist_directory=persist_directory, embedding_function=embeddings
    )
    if docs:
        texts = [d.page_content for d in docs]
        metadatas = [d.metadata for d in docs]
        vectorstore.add_texts(texts=texts, metadatas=metadatas)
        vectorstore.persist()
    logger.info("Loaded local resume, JD, and company infos.")
    # 사전 계산: 회사 정보와 체인 초기화
    stored_company_info = get_company_info()
    interview_chain = get_interview_chain()
    followup_chain = get_followup_chain()
    evaluate_chain = get_evaluate_chain()
    model_answer_chain = get_model_answer_chain()
    assessment_chain = get_assessment_chain()
    base_chain_inputs = {
        "resume": stored_resume,
        "jd": stored_jd,
        "company_infos": stored_company_info,
    }
    logger.info("Precomputed company_info and initialized chains.")


def get_company_info():
    if vectorstore is None:
        raise HTTPException(
            status_code=400,
            detail="Company documents not uploaded. Please upload docs first.",
        )
    # 회사 자료 검색
    retrieved = vectorstore.similarity_search(stored_jd, k=3)
    company_info = "\n".join([doc.page_content for doc in retrieved])
    # Trim company_info to avoid exceeding model context window
    max_company_info_length = 2000
    if len(company_info) > max_company_info_length:
        company_info = company_info[:max_company_info_length]
    logger.info(f"Retrieved company info length: {len(company_info)}")
    return company_info


@app.get("/chatHistory")
async def get_chat_history():
    return chat_history.get_all_history()


@app.get("/question")
async def generate_question():
    try:
        if interview_chain is None or base_chain_inputs is None:
            raise HTTPException(
                status_code=500, detail="Interview chain not initialized."
            )
        logger.info("Generating question using pre-initialized chain")
        response = interview_chain.invoke(base_chain_inputs)
        logger.info("Chain invocation completed")

        question_id = chat_history.add(
            type="question", speaker="agent", content=response["result"]
        )
        return {
            "id": question_id,
            "content": response["result"],
        }
    except Exception as e:
        logger.error(f"Error in question generation: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/answer")
async def submit_answer(request: AnswerRequest):
    question_id = request.questionId
    content = request.content
    if not chat_history.validate_question_exists(question_id):
        raise HTTPException(status_code=404, detail="Cannot find question id.")

    answer_id = chat_history.add(type="answer", speaker="user", content=content)
    """
    답변 평가 await 
    """
    evaluate_content = evaluate_chain.invoke(
        {
            **base_chain_inputs,
            "prev_question_answer_pairs": [
                {
                    "question": chat_history.get_question_by_id(question_id),
                    "answer": content,
                }
            ],
        }
    )

    chat_history.add(
        type="evaluate",
        speaker="agent",
        content=evaluate_content["result"],
    )
    return {"id": answer_id, "content": content}


@app.get("/assessment")
async def get_assessment():
    """
    평가 결과 반환
    논리성, 직무적합성, 핵심가치 부합성, 커뮤니케이션 능력
    """
    assessment_content = assessment_chain.invoke(
        {
            **base_chain_inputs,
            "chat_history": chat_history.get_all_history_as_string(),
        }
    )
    return assessment_content


@app.get("/followUp")
async def generate_followup(questionId: str):
    question_id = questionId
    if chat_history.history is None:
        raise HTTPException(status_code=400, detail="No chat history yet.")
    if not chat_history.validate_question_exists(question_id):
        raise HTTPException(status_code=404, detail="Cannot find question id.")

    try:
        logger.info("Generating followup using pre-initialized chain")
        if followup_chain is None or base_chain_inputs is None:
            raise HTTPException(
                status_code=500, detail="Followup chain not initialized."
            )

        question_item = None
        answer_item = None

        for idx, item in enumerate(chat_history.history):
            if item.id == question_id and item.type == "question":
                question_item = item
                if idx + 1 < len(chat_history.history):
                    next_item = chat_history.history[idx + 1]
                    if next_item.type == "answer":
                        answer_item = next_item
                break

        if question_item is None or answer_item is None:
            raise HTTPException(
                status_code=400,
                detail="No question or answer matching with questionId.",
            )

        inputs = {
            **base_chain_inputs,
            "prev_question_answer_pairs": [
                {"question": question_item, "answer": answer_item}
            ],
        }

        response = followup_chain.invoke(inputs)
        question_id = chat_history.add(
            type="question", speaker="agent", content=response["result"]
        )
        logger.info("Followup generated")
        return {
            "id": question_id,
            "content": response["result"],
        }
    except Exception as e:
        logger.error(f"Error in followup generation: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/modelAnswer")
async def generate_model_answer(questionId: str):
    question_id = questionId
    if chat_history.history is None:
        raise HTTPException(status_code=400, detail="No chat history yet.")
    if not chat_history.validate_question_exists(question_id):
        raise HTTPException(status_code=404, detail="Cannot find question id.")

    question_item = chat_history.get_question_by_id(question_id)
    if question_item is None:
        raise HTTPException(status_code=404, detail="Cannot find question data.")

    try:
        # TODO: response <- model_answer_chain 연결 필요
        response = model_answer_chain.invoke(
            {
                **base_chain_inputs,
                "question": question_item,
            }
        )
        answer_id = chat_history.add(
            type="modelAnswer", speaker="agent", content=response["result"]
        )
        return {"id": answer_id, "content": response["result"]}
    except Exception as e:
        logger.error(f"Error in model answer generation: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


app.mount("/", StaticFiles(directory=dist_path, html=True), name="frontend")

if __name__ == "__main__":
    uvicorn.run("app:app", host="0.0.0.0", port=8000, reload=True)
