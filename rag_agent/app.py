from fastapi import FastAPI, HTTPException
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field
from typing import Callable, Literal, Optional
import uvicorn
from rag_agent.chains.interview_chain import get_followup_chain, get_interview_chain
import logging
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import OpenAIEmbeddings
from langchain.text_splitter import CharacterTextSplitter
from langchain.docstore.document import Document
import os
import io
import shutil

# PDF/DOCX 파싱
import PyPDF2
import docx

# 로깅 설정
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(title="AI Interview Simulator")
dist_path = os.path.join(os.path.dirname(__file__), "../frontend/dist")

# 메모리 컨텍스트 저장 변수
stored_resume: Optional[str] = None
stored_jd: Optional[str] = None
# 사전 계산된 회사 정보 및 체인 저장 변수
stored_company_info: Optional[str] = None
interview_chain = None
followup_chain = None
base_chain_inputs: Optional[dict] = None
# RAG 벡터 스토어
vectorstore: Optional[Chroma] = None
# 영속 디렉토리 설정 (환경변수 또는 기본 경로)
persist_directory = os.getenv(
    "CHROMA_DB_PATH", os.path.join(os.path.dirname(__file__), "vectorstore/chroma_db")
)


class ChatHistory(BaseModel):
    """question, answer 쌍을 저장하는 클래스
    질문-답변 쌍으로 저장이 되어야하고
    질문-답변 쌍을 추가하는 메서드와
    질문-답변 쌍을 반환하는 메서드가 있어야한다.
    모든 질문-답변 쌍을 반환하는 메서드가 있어야 한다.
    """

    question_history: list[dict[str, str]] = Field(default_factory=list)
    answer_history: list[dict[str, str]] = Field(default_factory=list)

    def add_question(self, question: str):
        self.question_history.append(
            {"question_id": len(self.question_history), "question": question}
        )
        return self.question_history[-1]["question_id"]

    def add_answer(self, question_id: str, answer: str):
        self.answer_history.append({"question_id": question_id, "answer": answer})
        return self.answer_history[-1]["question_id"]

    def get_all_history(self) -> list[dict[str, str]]:
        """List[{question, answer}]"""
        return [
            {"question": q["question"], "answer": a["answer"]}
            for q, a in zip(self.question_history, self.answer_history)
        ]


chat_history = ChatHistory()


# 로컬 파일 시스템에서 context와 회사 자료 자동 로딩
def parse_file_to_text(file_path: str) -> str:
    with open(file_path, "rb") as f:
        content = f.read()
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        if file_path.lower().endswith(".pdf"):
            reader = PyPDF2.PdfReader(io.BytesIO(content))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        elif file_path.lower().endswith((".docx", ".doc")):
            doc = docx.Document(io.BytesIO(content))
            return "\n".join(p.text for p in doc.paragraphs)
        else:
            return content.decode("utf-8", errors="ignore")


async def init_local_data():
    """로컬 db를 모두 비운다"""
    global vectorstore
    if os.path.exists(persist_directory):
        shutil.rmtree(persist_directory)
    os.makedirs(persist_directory, exist_ok=True)
    vectorstore = None
    logger.info("Vectorstore reset successfully.")


@app.on_event("startup")
async def load_local_data():
    await init_local_data()
    global stored_resume, stored_jd, vectorstore, stored_company_info, interview_chain, followup_chain, base_chain_inputs, chat_history
    base_dir = os.path.join(os.path.dirname(__file__), "../data")
    # 이력서 로딩
    resume_dir = os.path.join(base_dir, "resume")
    for fname in os.listdir(resume_dir):
        stored_resume = parse_file_to_text(os.path.join(resume_dir, fname))
        break
    # JD 로딩
    jd_dir = os.path.join(base_dir, "jd")
    for fname in os.listdir(jd_dir):
        stored_jd = parse_file_to_text(os.path.join(jd_dir, fname))
        break
    # 회사 자료 로딩 및 인덱싱
    company_dir = os.path.join(base_dir, "company_infos")
    docs = []
    for fname in os.listdir(company_dir):
        text = parse_file_to_text(os.path.join(company_dir, fname))
        splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
        for chunk in splitter.split_text(text):
            docs.append(Document(page_content=chunk, metadata={"filename": fname}))
    embeddings = OpenAIEmbeddings()
    vectorstore = Chroma(
        persist_directory=persist_directory, embedding_function=embeddings
    )
    if docs:
        texts = [d.page_content for d in docs]
        metadatas = [d.metadata for d in docs]
        vectorstore.add_texts(texts=texts, metadatas=metadatas)
        vectorstore.persist()
    logger.info("Loaded local resume, JD, and company infos.")
    # 사전 계산: 회사 정보와 체인 초기화
    stored_company_info = get_company_info()
    interview_chain = get_interview_chain()
    followup_chain = get_followup_chain()
    base_chain_inputs = {
        "resume": stored_resume,
        "jd": stored_jd,
        "company_info": stored_company_info,
    }
    logger.info("Precomputed company_info and initialized chains.")


class AnswerRequest(BaseModel):
    question_id: str
    user_answer: str


class FollowUpRequest(BaseModel):
    user_answer: str
    mode: str  # "tail", "model", "next"


def get_company_info():
    if vectorstore is None:
        raise HTTPException(
            status_code=400,
            detail="Company documents not uploaded. Please upload docs first.",
        )
    # 회사 자료 검색
    retrieved = vectorstore.similarity_search(stored_jd, k=3)
    company_info = "\n".join([doc.page_content for doc in retrieved])
    # Trim company_info to avoid exceeding model context window
    max_company_info_length = 2000
    if len(company_info) > max_company_info_length:
        company_info = company_info[:max_company_info_length]
    logger.info(f"Retrieved company info length: {len(company_info)}")
    return company_info


@app.post("/question")
async def generate_question():
    try:
        if interview_chain is None or base_chain_inputs is None:
            raise HTTPException(
                status_code=500, detail="Interview chain not initialized."
            )
        logger.info("Generating question using pre-initialized chain")
        response = interview_chain.invoke(base_chain_inputs)
        logger.info("Chain invocation completed")

        question_id = chat_history.add_question(response["result"])
        return {
            "question": {
                "question_id": question_id,
                "question": response["result"],
            }
        }
    except Exception as e:
        logger.error(f"Error in question generation: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/evaluate")
async def evaluate_answer(request: AnswerRequest):
    print(request)
    question_id = request.question_id
    answer = request.user_answer
    chat_history.add_answer(question_id, answer)
    return {"response": answer}


@app.post("/followup")
async def generate_followup():
    if chat_history.question_history is None:
        raise HTTPException(status_code=400, detail="No question generated yet.")
    if chat_history.answer_history is None:
        raise HTTPException(status_code=400, detail="No answer generated yet.")
    try:
        logger.info("Generating followup using pre-initialized chain")
        if followup_chain is None or base_chain_inputs is None:
            raise HTTPException(
                status_code=500, detail="Followup chain not initialized."
            )
        inputs = {
            **base_chain_inputs,
            "prev_question_answer_pairs": chat_history.get_all_history(),
        }
        response = followup_chain.invoke(inputs)
        question_id = chat_history.add_question(response["result"])
        logger.info("Followup generated")
        return {"response": question_id}
    except Exception as e:
        logger.error(f"Error in followup generation: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


app.mount("/", StaticFiles(directory=dist_path, html=True), name="frontend")

if __name__ == "__main__":
    uvicorn.run("app:app", host="0.0.0.0", port=8000, reload=True)
